import sys
from docx import Document

def update_template(input_file, output_file):
    doc = Document(input_file)
    
    replacements = {
        "En la ciudad de __________________, a ______________ días del mes de _______ del 202___": "En la ciudad de {city}, a {day_number_text} días del mes de {month_text} del {year_text}",
        "el(la) señor(a)(ita) ______________________________": "el(la) señor(a)(ita) {name_worker}",
        "cédula de ciudadanía/identidad/pasaporte No. ____________________": "cédula de ciudadanía/identidad/pasaporte No. {id_worker}",
        "domiciliada en el cantón Samborondón": "domiciliada en el cantón {city}",
        "cargo de INSTRUCTOR CAPACITADOR": "cargo de {work_title}",
        "De lunes a viernes, el(la) COLABORADOR(A) laborará siete horas diarias, en el horario de ____H____ a ____H____, incluyendo _____________ (_______) minutos": "De lunes a viernes, el(la) COLABORADOR(A) laborará siete horas diarias, en el horario de {start_time_weekday} a {end_time_weekday}, incluyendo {break_minutes_weekday_text} ({break_minutes_weekday}) minutos",
        "Los días sábados, el(la) TRABAJADOR(A) laborará en el horario de ____H____ a ____H____, incluyendo treinta (30) minutos": "Los días sábados, el(la) TRABAJADOR(A) laborará en el horario de {start_time_saturday} a {end_time_saturday}, incluyendo treinta (30) minutos",
        "ubicado en el Centro Comercial “Samborondón Plaza”, planta baja, local No. 13, Km. 1.5 vía a Samborondón, provincia del Guayas, en la ciudad de Samborondón": "ubicado en {city_working_place}, en la ciudad de {city}",
        "domicilio está ubicado en ____________________________________________ (_____________________________), en la ciudad de __________________, provincia de _______________": "domicilio está ubicado en {domicile_worker} ({domicile_reference}), en la ciudad de {city_name_worker}, provincia de {province_name_worker}",
        "correo electrónico personal es “_______________@_________._____”": "correo electrónico personal es \"{mail_worker}\"",
        "unificada de CUATROCIENTOS NOVENTA Y SEIS 45/100 DÓLARES DE LOS ESTADOS UNIDOS DE AMÉRICA (US$496.45)": "unificada de {salary_text} {salary_cents_text} DÓLARES DE LOS ESTADOS UNIDOS DE AMÉRICA (US${salary_value})",
        "jueces multicompetentes del cantón Samborondón, provincia del Guayas": "jueces multicompetentes del cantón {city}, provincia de {province}",
        "HUGO G. LOPEZ GOMEZ    ____________________________": "HUGO G. LOPEZ GOMEZ    {name_worker}",
        "PRESIDENTE    C.C. No. ___________________": "PRESIDENTE    C.C. No. {id_worker}",
        "fuera de la provincia de Guayas": "fuera de la provincia de {province}",
        "cargo de VENDEDOR(A)-PROMOTOR(A)": "cargo de {work_title}",
        "unificada de CUATROCIENTOS NOVENTA Y TRES 90/100 DÓLARES DE LOS ESTADOS UNIDOS DE AMÉRICA (US$493.90)": "unificada de {salary_text} {salary_cents_text} DÓLARES DE LOS ESTADOS UNIDOS DE AMÉRICA (US${salary_value})"
    }

    # Handling the functions list
    in_functions = False
    paragraphs_to_remove = []
    function_insert_index = -1
    item_style = None
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # Replace normal texts
        for old_str, new_str in replacements.items():
            if old_str in p.text:
                replaced = False
                for run in p.runs:
                    if old_str in run.text:
                        run.text = run.text.replace(old_str, new_str)
                        replaced = True
                
                if not replaced:
                    style = p.style
                    new_text = p.text.replace(old_str, new_str)
                    p.clear()
                    p.add_run(new_text)
                    p.style = style

        if "En virtud de dicho cargo el(la) EMPLEADO(A) deberá ejecutar las labores o funciones siguientes:" in text:
            in_functions = True
            function_insert_index = i + 1
            continue
            
        if in_functions:
            if "CLÁUSULA TERCERA" in text:
                in_functions = False
            else:
                if item_style is None and p.style is not None and text != "":
                    item_style = p.style
                paragraphs_to_remove.append(p)
                
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)
        
    if function_insert_index != -1:
        target_p = doc.paragraphs[function_insert_index]
        
        p_start = target_p.insert_paragraph_before('{#funciones_lista}')
        p_item = target_p.insert_paragraph_before('{item}')
        if item_style:
            try:
                p_item.style = item_style
            except:
                pass
        p_end = target_p.insert_paragraph_before('{/funciones_lista}')

    doc.save(output_file)
    print(f"Template successfully saved to {output_file}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python update_docx.py <input> <output>")
        sys.exit(1)
    update_template(sys.argv[1], sys.argv[2])
